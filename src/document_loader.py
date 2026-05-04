"""Document ingestion — scans /documents and extracts clean text."""
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class Document:
    content: str
    source: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def _load_txt(path: Path) -> Optional[str]:
    try:
        return path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as exc:
        logger.warning("TXT read error %s: %s", path.name, exc)
        return None


def _load_markdown(path: Path) -> Optional[str]:
    return _load_txt(path)


def _load_pdf(path: Path) -> Optional[str]:
    try:
        import pymupdf  # fitz

        doc = pymupdf.open(str(path))
        pages = [page.get_text() for page in doc]
        doc.close()
        text = "\n".join(pages).strip()
        return text if text else None
    except ImportError:
        logger.error("pymupdf not installed. Run: pip install pymupdf")
        return None
    except Exception as exc:
        logger.warning("PDF read error %s: %s", path.name, exc)
        return None


def _load_docx(path: Path) -> Optional[str]:
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text.strip() if text.strip() else None
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as exc:
        logger.warning("DOCX read error %s: %s", path.name, exc)
        return None


_LOADERS = {
    ".txt": _load_txt,
    ".md": _load_markdown,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


def load_documents(documents_dir: Path) -> List[Document]:
    """Recursively scan *documents_dir* and return a list of Document objects."""
    docs: List[Document] = []

    if not documents_dir.exists():
        logger.warning("Documents directory not found: %s", documents_dir)
        return docs

    files = [f for f in documents_dir.rglob("*") if f.is_file()]
    logger.info("Found %d file(s) in %s", len(files), documents_dir)

    for file_path in sorted(files):
        ext = file_path.suffix.lower()
        loader = _LOADERS.get(ext)

        if loader is None:
            logger.debug("Unsupported file type, skipping: %s", file_path.name)
            continue

        logger.info("Loading %s ...", file_path.name)
        text = loader(file_path)

        if not text:
            logger.warning("Empty or unreadable file, skipping: %s", file_path.name)
            continue

        docs.append(
            Document(
                content=text,
                source=file_path.name,
                file_type=ext.lstrip("."),
                metadata={"path": str(file_path), "size_bytes": file_path.stat().st_size},
            )
        )

    logger.info("Successfully loaded %d document(s)", len(docs))
    return docs
